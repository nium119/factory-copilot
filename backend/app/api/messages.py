"""
消息API
提供消息发送接口，支持 Agent 路由
"""
import json
from typing import Any, Dict, List, Optional

from fastapi import APIRouter, Depends, HTTPException, Request
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine

from app.agents import get_agents_from_db
from app.core.config import settings
from app.core.logger import log
from app.repositories.conversation_repository import ConversationRepository
from app.repositories.message_repository import MessageRepository
from app.services.conversation_service import ConversationService
from app.services.message_service import MessageService

router = APIRouter(prefix="/messages", tags=["消息"])


@router.get("/events/stream")
async def event_stream():
    """全局 SSE 事件流 — 审批状态变更实时推送。"""
    from app.services.event_bus import event_bus
    return StreamingResponse(
        event_bus.subscribe(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
    )


class SendMessageRequest(BaseModel):
    """发送消息请求"""
    conversation_id: str
    content: str
    model_name: Optional[str] = None
    agent_name: Optional[str] = None
    use_agent: bool = False
    web_search: bool = False
    enable_memory: bool = True
    enable_thinking: Optional[bool] = None


class RestoreEntityRequest(BaseModel):
    """轨迹回滚：恢复被删实体（delete）或删除新建实体（create）。"""
    tool: str
    records: List[Dict[str, Any]] = []
    created_entity_id: Optional[str] = None


# 模块级引擎和会话工厂，应用启动时创建一次

_engine = create_async_engine(settings.DATABASE_URL, echo=False)
_async_session = async_sessionmaker(_engine, expire_on_commit=False)


async def get_db() -> AsyncSession:
    """获取数据库会话（复用全局引擎）"""
    async with _async_session() as session:
        yield session


def get_message_service(db: AsyncSession = Depends(get_db)) -> MessageService:
    """获取消息服务实例"""
    message_repo = MessageRepository(db)
    conversation_repo = ConversationRepository(db)
    return MessageService(message_repo, conversation_repo)


def get_conversation_service(db: AsyncSession = Depends(get_db)) -> ConversationService:
    """获取会话服务实例"""
    conversation_repo = ConversationRepository(db)
    message_repo = MessageRepository(db)
    return ConversationService(conversation_repo, message_repo)


def get_current_user_id(request: Request) -> str:
    """从 Bearer token 验签解析当前用户。

    安全修复：移除 X-User-Id 直传（客户端可伪造身份），统一 Bearer JWT 验签。
    """
    from app.api.deps import get_current_user_id as _resolve
    return _resolve(request)


class ConfirmRequest(BaseModel):
    """确认请求"""
    approved: bool = True
    params: Optional[dict] = None


class PendingQuery(BaseModel):
    """待办查询参数"""
    user_id: str = ""
    user_roles: str = ""  # 逗号分隔的角色列表


class ApprovalRequest(BaseModel):
    """审批请求"""
    user_id: str = ""
    comment: str = ""


@router.post("/confirm/{session_id}", summary="确认或取消写操作")
async def confirm_action(session_id: str, request: ConfirmRequest):
    """前端确认/取消高危操作。

    后端 _standard_process 在执行 requiresConfirmation 的 Action 前
    会挂起等待此端点。超时 60s 自动取消。
    """
    from app.agents.base import BaseAgent
    # 先读消息和操作名（resolve 会 pop 掉数据）
    entry = BaseAgent._pending_confirmations.get(session_id, {})
    action_label = entry.get("action_label", "") if isinstance(entry, dict) else ""
    user_message = entry.get("message", "") if isinstance(entry, dict) else ""
    resolved = BaseAgent.resolve_confirmation(session_id, request.approved, request.params)
    # 取消时记录负反馈
    if not request.approved and action_label:
        from app.db import get_db
        from app.models.intent_feedback import IntentFeedback
        async for db_session in get_db():
            db_session.add(IntentFeedback(
                message=user_message,
                matched_action=action_label,
                was_correct=0,
            ))
            await db_session.commit()
            break
    return {"resolved": resolved, "session_id": session_id, "approved": request.approved}


@router.get("/agents", summary="获取可用 Agent 列表")
async def list_agents():
    """从数据库获取所有已注册的 Agent 元信息"""
    from app.db import get_db
    from app.repositories.agent_repository import AgentRepository
    async for session in get_db():
        repo = AgentRepository(session)
        agents = await repo.get_enabled_agents()
        result = []
        for a in agents:
            try:
                from app.agents import get_agent
                agent = get_agent(a.name)
                info = agent.get_info()
                info["enabled"] = a.enabled
                result.append(info)
            except (KeyError, Exception):
                result.append({
                    "name": a.name, "display_name": a.display_name,
                    "icon": a.icon, "color": a.color,
                    "description": a.description, "enabled": a.enabled,
                })
        return result


@router.post("/stream", summary="流式发送消息")
async def send_message_stream(
    request: SendMessageRequest,
    http_request: Request,
    message_service: MessageService = Depends(get_message_service),
    user_id: str = Depends(get_current_user_id)
):
    """
    通过 SSE 流式发送消息并逐步接收 AI 回复。

    返回的 SSE 事件类型：
    - **agent_info**: Agent 元信息（首次发送）
    - **content**: AI 回复的文本片段
    - **thinking**: 模型思考过程
    - **error**: 错误信息
    - **[DONE]**: 传输完成标记
    """

    # 从请求头提取 MES token，传递给 CLI 子进程
    mes_token = None
    auth_header = http_request.headers.get("Authorization", "")
    if auth_header.startswith("Bearer "):
        mes_token = auth_header[7:]

    async def event_generator():
        """SSE事件生成器"""
        # 设置 API 调用日志的请求上下文
        from app.services.multi_system_backend import _request_user_id, _request_conversation_id, _request_message
        _request_user_id.set(user_id or "")
        _request_conversation_id.set(request.conversation_id)
        _request_message.set(request.content[:200])
        try:
            from app.agents import get_agent
            from app.agents.router import route_intent
            from app.agents.tools.mes_cli_runner import set_token

            set_token(mes_token)

            route = await route_intent(request.content, request.agent_name)
            agent_name = route["agent_name"]
            use_agent = route["use_agent"]
            matched_agents = route.get("matched_agents", [])

            # 用户偏好适应
            if route["confidence"] < 0.9:
                try:
                    async with _async_session() as adapt_db:
                        from app.services.adaptation_service import get_adapted_confidence
                        adapted_confidence = await get_adapted_confidence(
                            adapt_db, user_id, agent_name, route["confidence"]
                        )
                        if adapted_confidence < 0.3:
                            log.info(
                                f"[Adaptation] 用户偏好负向，{agent_name} 置信度过低 ({adapted_confidence})，"
                                f"回退到 analysis_monitor"
                            )
                            agent_name = "analysis_monitor"
                            use_agent = False
                except Exception as e:
                    log.debug(f"[Adaptation] 偏好调整跳过: {e}")

            # 资源感知：根据查询复杂度自动选择模型
            from app.agents.router import select_model_for_complexity
            model_name = select_model_for_complexity(request.content, request.model_name)
            if model_name:
                log.info(f"[资源优化] 自动选择模型: {model_name} (基于查询复杂度)")
            else:
                model_name = request.model_name

            try:
                agent = get_agent(agent_name)
            except KeyError:
                # Agent 不可用时，先尝试链引擎匹配
                from app.core.chain_engine import chain_engine, _CHAINS, reload_chains_async
                if not _CHAINS:
                    try: await reload_chains_async()
                    except: pass
                chain_id = await chain_engine.detect(request.content)
                if chain_id:
                    log.info(f"[SSE] Agent不可用但链匹配: {chain_id}")
                    chain_engine.set_agent_resolver(get_agent)
                    async for cht, chc in chain_engine.execute(
                        message=request.content, chain_id=chain_id,
                        model_name=model_name, enable_thinking=request.enable_thinking,
                        session_id=request.conversation_id,
                    ):
                        yield f"data: {json.dumps({'type': cht, 'content': chc})}\n\n"
                    yield "data: [DONE]\n\n"
                    return
                else:
                    log.warning(f"[SSE] Agent '{agent_name}' 不可用（无域配置）")
                    yield f"data: {json.dumps({'type': 'error', 'content': '当前没有可用 Agent，请先配置业务域'})}\n\n"
                    yield "data: [DONE]\n\n"
                    return
            agent_info = agent.get_info()

            # 发送 Agent 信息
            log.info(f"[SSE] 发送 agent_info: {agent_info['display_name']}")
            yield f"data: {json.dumps({'type': 'agent_info', 'agent_name': agent_info['name'], 'display_name': agent_info['display_name'], 'icon': agent_info['icon'], 'color': agent_info['color']})}\n\n"

            # 通过 MessageService 处理消息（包含记忆、数据库持久化）
            # 传递已路由的 agent_name，避免 service 层重复路由（LLM 调用导致延迟）
            async for chunk_type, chunk_content in message_service.process_message_stream(
                user_id=user_id,
                conversation_id=request.conversation_id,
                message=request.content,
                model_name=model_name,
                use_agent=use_agent,
                web_search=request.web_search,
                enable_memory=request.enable_memory,
                agent_name=agent_name,
                enable_thinking=request.enable_thinking,
                matched_agents=matched_agents,
            ):
                log.info(f"[SSE] 发送 chunk_type={chunk_type}, content_len={len(str(chunk_content))}")
                yield f"data: {json.dumps({'type': chunk_type, 'content': chunk_content})}\n\n"

            # 发送数据来源信息 + 恢复建议
            from app.agents.tools.mes_cli_runner import get_data_source, get_error_recovery_hint
            ds = get_data_source()
            hint = get_error_recovery_hint() if ds == "mock_fallback" else None
            yield f"data: {json.dumps({'type': 'data_source', 'source': ds, 'hint': hint})}\n\n"

            # 发送结束标记
            log.info("[SSE] 发送 [DONE]")
            yield "data: [DONE]\n\n"

        except Exception as e:
            error_data = json.dumps({
                "type": "error",
                "content": str(e)
            })
            yield f"data: {error_data}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        }
    )


@router.get("/pending")
async def get_pending_confirmations(
    user_id: str = "",
    user_roles: str = "",
    page: int = 1,
    page_size: int = 20,
    db: AsyncSession = Depends(get_db),
):
    """获取当前用户可审批的待办消息列表（分页）。"""
    from app.repositories.message_repository import MessageRepository
    repo = MessageRepository(db)

    roles = [r.strip() for r in user_roles.split(",") if r.strip()] if user_roles else []
    offset = max(0, (page - 1)) * max(1, page_size)

    # 查询待审批消息（拉取较大窗口以支持角色过滤后仍有足够结果）
    all_pending = await repo.get_pending_confirmations(
        assigned_to=None, limit=max(page_size * 3, 100), offset=0,
    )
    total_count = await repo.count_pending_confirmations(assigned_to=None)

    # 按角色过滤
    filtered = []
    for msg in all_pending:
        assigned = msg.assigned_to or ""
        if not roles or not assigned or assigned in roles:
            content_data = {}
            try:
                content_data = json.loads(msg.content) if msg.content else {}
            except Exception:
                content_data = {"raw": msg.content}
            filtered.append({
                "id": msg.id,
                "conversation_id": msg.conversation_id,
                "conversation_title": content_data.get("conversation_title", ""),
                "message_type": msg.message_type or "",
                "action_label": content_data.get("action_label", ""),
                "concept_label": content_data.get("concept_label", ""),
                "tool": content_data.get("tool", ""),
                "params": content_data.get("params", {}),
                "param_schema": content_data.get("param_schema", []),
                "risk": content_data.get("risk", "write"),
                "context": content_data.get("context", {}),
                "decision_pack": content_data.get("decision_pack", {}),
                "user_id": content_data.get("user_id", ""),
                "message": content_data.get("message", ""),
                "error_detail": content_data.get("error_detail", ""),
                "verify_detail": content_data.get("verify_detail", []),
                "verify_summary": content_data.get("verify_summary", ""),
                "submitter_id": content_data.get("submitter_id", ""),
                "chain_id": content_data.get("chain_id", ""),
                "message_id": content_data.get("message_id", ""),
                "assigned_to": assigned,
                "created_at": str(msg.created_at) if msg.created_at else "",
            })

    # 审批条目（confirm）content 无 conversation_title：从会话表回退查，保证「📎 原对话」入口可用
    _missing = [it for it in filtered if not it.get("conversation_title") and it.get("conversation_id")]
    if _missing:
        from app.repositories.conversation_repository import ConversationRepository
        _cr = ConversationRepository(db)
        for _it in _missing:
            try:
                _cv = await _cr.get_by_id(_it["conversation_id"])
                if _cv and _cv.title:
                    _it["conversation_title"] = _cv.title
            except Exception:
                pass

    # 分页切片
    paged = filtered[offset:offset + max(1, page_size)]
    return {
        "pending": paged, "total": len(filtered),
        "page": page, "page_size": page_size,
        "have_param_schema": True,
    }


@router.get("/reports")
async def get_reports(page: int = 1, page_size: int = 20, db: AsyncSession = Depends(get_db)):
    """获取历史分析报告（message_type=report 的消息），按时间倒序，分页。"""
    from app.repositories.message_repository import MessageRepository
    from app.models.message import Message
    from sqlalchemy import select, func
    import json
    repo = MessageRepository(db)
    q = select(Message).where(
        Message.message_type == "report"
    ).order_by(Message.created_at.desc()).offset((page - 1) * page_size).limit(page_size)
    r = await db.execute(q)
    items = list(r.scalars().all())
    result = []
    for msg in items:
        # 查会话标题
        conv_title = ""
        user_question = ""
        try:
            from app.repositories.conversation_repository import ConversationRepository
            conv_repo = ConversationRepository(db)
            conv = await conv_repo.get_by_id(msg.conversation_id)
            if conv:
                conv_title = conv.title or ""
        except Exception:
            pass
        result.append({
            "id": msg.id,
            "conversation_id": msg.conversation_id,
            "content": msg.content,
            "title": conv_title,
            "metadata": msg.metadata_dict or {},
            "created_at": str(msg.created_at) if msg.created_at else "",
        })
    # 总数
    cq = select(func.count()).where(Message.message_type == "report")
    cr = await db.execute(cq)
    total = cr.scalar() or 0
    return {"reports": result, "total": total, "page": page, "page_size": page_size}


@router.get("/reports/{report_id}/export")
async def export_report(
    report_id: str,
    format: str = "pdf",
    db: AsyncSession = Depends(get_db),
):
    """导出报告为 PDF 或 Word 格式。

    GET /api/messages/reports/{id}/export?format=pdf   → application/pdf
    GET /api/messages/reports/{id}/export?format=docx  → Word 文档
    """
    from fastapi.responses import Response
    from app.repositories.message_repository import MessageRepository

    repo = MessageRepository(db)
    msg = await repo.get_by_id(report_id)
    if not msg:
        raise HTTPException(status_code=404, detail="报告不存在")

    md_content = msg.content or ""
    # 隐藏变更方案 JSON 代码块
    import re as _re
    md_content = _re.sub(r'```(?:json)?\s*\n[\s\S]*?\n```', '', md_content)
    # 追加变更方案详情到导出内容
    meta = msg.metadata_dict or {}
    change_plans = meta.get("change_plans", [])
    if change_plans:
        lines = ["\n\n---\n## 📋 变更方案\n"]
        for i, plan in enumerate(change_plans):
            risk_map = {"low": "🟢 低风险", "medium": "🟡 中风险", "high": "🔴 高风险"}
            lines.append(f"### 方案{i + 1}：{plan.get('label', '')}")
            lines.append(f"- **前提条件**：{plan.get('precondition', '')}")
            lines.append(f"- **影响评估**：{plan.get('impact', '')}")
            lines.append(f"- **风险等级**：{risk_map.get(plan.get('risk', ''), plan.get('risk', ''))}")
            steps = plan.get("steps_preview", [])
            if steps:
                lines.append("- **执行步骤**：")
                for j, s in enumerate(steps):
                    lines.append(f"  {j + 1}. {s}")
            lines.append("")
        md_content += "\n".join(lines)
    title = ""
    try:
        from app.repositories.conversation_repository import ConversationRepository
        conv_repo = ConversationRepository(db)
        conv = await conv_repo.get_by_id(msg.conversation_id)
        if conv:
            title = conv.title or "分析报告"
    except Exception:
        title = "分析报告"

    if format == "pdf":
        # PDF：生成自包含 HTML 页面，CDN 加载 ECharts/Mermaid 在浏览器端渲染真实图表后打印
        html = _build_print_html(md_content, title)
        return Response(content=html, media_type="text/html; charset=utf-8")

    elif format == "docx":
        # Word：ECharts → 原生 Office 图表，Mermaid → mermaid.ink SVG
        try:
            docx_bytes = _build_docx_with_charts(md_content, title)
            filename = f"{title}.docx"
            return Response(content=docx_bytes,
                            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{_url_quote(filename)}"})
        except Exception as e:
            log.warning(f"Word 生成失败: {e}")
            raise HTTPException(status_code=500, detail=f"Word 生成失败: {e}")

    else:
        # HTML 预览
        html_body = _md_to_docx_html(md_content)
        html = _build_print_html(md_content, title).replace(
            '<script>window.addEventListener', '<!--')
        return Response(content=html, media_type="text/html; charset=utf-8")


def _build_print_html(md_text: str, title: str) -> str:
    """生成自包含 HTML 页面，CDN 加载 ECharts/Mermaid 在浏览器端渲染真实图表后自动打印。"""
    import re as _re
    import json as _json
    import base64

    # 提取所有 echarts/mermaid 块的索引和内容，替换为容器 div
    echarts_blocks = []
    mermaid_blocks = []

    def _collect_echarts(match):
        code = match.group(1).strip()
        idx = len(echarts_blocks)
        echarts_blocks.append(code)
        return f'<div id="echarts-{idx}" class="echarts-chart" style="width:100%;min-height:350px;margin:16px 0;"></div>'

    def _collect_mermaid(match):
        code = match.group(1).strip()
        idx = len(mermaid_blocks)
        mermaid_blocks.append(code)
        return f'<div class="mermaid" style="margin:16px 0;">{code}</div>'

    processed = _re.sub(r'```echarts\s*\n(.*?)\n```', _collect_echarts, md_text, flags=_re.DOTALL)
    processed = _re.sub(r'```mermaid\s*\n(.*?)\n```', _collect_mermaid, processed, flags=_re.DOTALL)

    # Markdown → HTML
    import markdown as _md
    html_body = _md.markdown(processed, extensions=['tables', 'fenced_code', 'codehilite', 'toc'])

    # 构建 echarts 初始化 JS
    echarts_js = ""
    if echarts_blocks:
        opts_json = []
        for i, code in enumerate(echarts_blocks):
            opt = _parse_echarts_option(code)
            opts_json.append(opt)
        echarts_js = f"""
<script>
let chartOpts = {_json.dumps(opts_json, ensure_ascii=False)};
let chartsReady = 0;
let totalCharts = chartOpts.length;
function initEcharts() {{
    chartOpts.forEach(function(opt, i) {{
        let el = document.getElementById('echarts-' + i);
        if (el && opt && Object.keys(opt).length > 0) {{
            let chart = echarts.init(el);
            chart.setOption(opt);
            chartsReady++;
        }}
    }});
    tryAutoPrint();
}}
</script>"""

    # 构建 mermaid JS
    mermaid_js = ""
    if mermaid_blocks:
        mermaid_js = """
<script>
let mermaidReady = false;
mermaid.initialize({{ startOnLoad: true, theme: 'default', securityLevel: 'loose' }});
</script>"""

    # 自动打印逻辑：先调 initEcharts 渲染图表，等 echarts + mermaid 完成后触发打印
    auto_print_js = f"""
<script>
let echartsDone = {str(not echarts_blocks).lower()};
let mermaidDone = {str(not mermaid_blocks).lower()};
function tryAutoPrint() {{
    if (echartsDone && mermaidDone) {{
        setTimeout(function() {{ window.print(); }}, 600);
    }}
}}
// 初始化图表并设置完成回调
document.addEventListener('DOMContentLoaded', function() {{
    if (typeof initEcharts === 'function') {{
        initEcharts();
        echartsDone = true;
    }}
    tryAutoPrint();
}});
// Mermaid polls async completion
let mermaidCheck = setInterval(function() {{
    let svgs = document.querySelectorAll('.mermaid svg');
    if (svgs.length >= {len(mermaid_blocks)}) {{
        mermaidDone = true;
        clearInterval(mermaidCheck);
        tryAutoPrint();
    }}
}}, 200);
// Fallback: auto-print after 10s regardless
setTimeout(function() {{
    echartsDone = true; mermaidDone = true;
    tryAutoPrint();
}}, 10000);
</script>"""

    # CSS 样式表
    css = """
body { font-family: "Microsoft YaHei", "SimSun", sans-serif; font-size: 14px; line-height: 1.8;
       max-width: 960px; margin: 30px auto; padding: 0 24px; color: #333; }
h1 { font-size: 1.8em; border-bottom: 2px solid #2563eb; padding-bottom: 8px; color: #1e3a5f; }
h2 { font-size: 1.4em; border-bottom: 1px solid #e0e0e0; padding-bottom: 6px; color: #333; margin-top: 28px; }
h3 { font-size: 1.15em; color: #444; }
table { border-collapse: collapse; width: 100%; margin: 12px 0; font-size: 13px; }
th, td { border: 1px solid #d0d0d0; padding: 6px 10px; text-align: left; }
th { background: #e8ecf1; font-weight: 600; }
tr:nth-child(even) { background: #f8f9fb; }
pre { background: #f4f5f7; padding: 12px; border-radius: 4px; overflow-x: auto; font-size: 13px; }
code { background: #f0f0f0; padding: 1px 4px; border-radius: 2px; font-size: 0.9em; }
blockquote { border-left: 3px solid #6c5ce7; margin: 12px 0; padding: 4px 16px; color: #555; background: #f8f7ff; }
img { max-width: 100%; }
.echarts-chart { border: 1px solid #e8e8e8; border-radius: 8px; background: #fafafa; }
@media print {
    body { max-width: 100%; margin: 0; padding: 0 12px; -webkit-print-color-adjust: exact; print-color-adjust: exact; }
    .echarts-chart { break-inside: avoid; }
    h1, h2, h3 { break-after: avoid; }
    table { break-inside: avoid; }
}"""

    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title}</title>
<script src="https://cdn.jsdelivr.net/npm/echarts@5.5.0/dist/echarts.min.js"></script>
<script src="https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js"></script>
<style>{css}</style>
</head>
<body>
<h1>{title}</h1>
{html_body}
{echarts_js}
{mermaid_js}
{auto_print_js}
</body>
</html>"""


def _md_to_docx_html(md_text: str) -> str:
    """Markdown → HTML（DOCX 导出用：echarts → 表格，mermaid → 占位图片）。"""
    return _md_to_html(md_text)


def _parse_echarts_option(js_text: str) -> dict:
    """将 ECharts JS option 对象转为 Python dict（简易解析器）。"""
    import json as _json
    text = js_text.strip()
    # 先尝试直接 JSON 解析
    try:
        return _json.loads(text)
    except Exception:
        pass
    # 尝试 JS 对象字面量转 JSON
    text = _re.sub(r'^\s*option\s*=\s*', '', text)
    text = _re.sub(r';\s*$', '', text)
    text = _re.sub(r'(\s)(\w+)(\s*:)', r'\1"\2"\3', text)
    text = _re.sub(r'^(\w+)(\s*:)', r'"\1"\2', text)
    text = text.replace("'", '"')
    text = _re.sub(r',(\s*[}\]])', r'\1', text)
    try:
        return _json.loads(text)
    except Exception:
        return {}


def _render_echarts_blocks(md_text: str) -> str:
    """将 ```echarts 代码块渲染为 HTML 表格。"""
    import re as _re

    def _render_echarts(match):
        code = match.group(1)
        opt = _parse_echarts_option(code)
        if not opt:
            return f'<pre><code>{code}</code></pre>'

        title = ""
        if "title" in opt and isinstance(opt["title"], dict):
            title = opt["title"].get("text", "")
        elif "title" in opt and isinstance(opt["title"], str):
            title = opt["title"]

        # 提取 xAxis 类别
        categories = []
        xaxis = opt.get("xAxis", {})
        if isinstance(xaxis, dict):
            categories = xaxis.get("data", [])
        elif isinstance(xaxis, list) and len(xaxis) > 0:
            categories = xaxis[0].get("data", []) if isinstance(xaxis[0], dict) else []

        # 提取 series
        series_list = opt.get("series", [])
        if isinstance(series_list, dict):
            series_list = [series_list]

        # 构建 HTML 表格
        rows = []
        if title:
            rows.append(f'<tr><th colspan="{max(2, len(series_list) + 1)}" style="text-align:center;background:#f0f0f0;">{title}</th></tr>')

        if categories:
            header = "<tr><th></th>"
            for s in series_list:
                header += f'<th>{s.get("name", "")}</th>'
            header += "</tr>"
            rows.append(header)

            max_len = len(categories)
            for i in range(max_len):
                row = f"<td>{categories[i] if i < len(categories) else ''}</td>"
                for s in series_list:
                    data = s.get("data", [])
                    val = data[i] if i < len(data) else ""
                    row += f"<td>{val}</td>"
                rows.append(f"<tr>{row}</tr>")
        else:
            # 无类别的简单序列
            header = "<tr>"
            for s in series_list:
                header += f'<th>{s.get("name", "")}</th>'
            header += "</tr>"
            rows.append(header)
            max_len = max((len(s.get("data", [])) for s in series_list), default=0)
            for i in range(max_len):
                row = ""
                for s in series_list:
                    data = s.get("data", [])
                    row += f"<td>{data[i] if i < len(data) else ''}</td>"
                rows.append(f"<tr>{row}</tr>")

        chart_type = series_list[0].get("type", "chart") if series_list else "chart"
        return (
            f'<div style="margin:16px 0;padding:12px;border:1px solid #e0e0e0;border-radius:6px;background:#fafafa;">'
            f'<div style="font-size:12px;color:#888;margin-bottom:8px;">📊 {chart_type} 图表</div>'
            f'<table style="width:100%;font-size:13px;">{"".join(rows)}</table></div>'
        )

    return _re.sub(r'```echarts\s*\n(.*?)\n```', _render_echarts, md_text, flags=_re.DOTALL)


def _render_mermaid_blocks(md_text: str) -> str:
    """将 ```mermaid 代码块渲染为 mermaid.ink 图片。"""
    import base64, re as _re, zlib

    def _render_mermaid(match):
        code = match.group(1).strip()
        # 使用 mermaid.ink API: pako.deflate → base64url
        compressed = zlib.compress(code.encode("utf-8"))[2:-4]  # 去掉 zlib header/trailer，等同 raw deflate
        encoded = base64.urlsafe_b64encode(compressed).decode("ascii").rstrip("=")
        url = f"https://mermaid.ink/svg/{encoded}"
        return (
            f'<div style="margin:16px 0;text-align:center;">'
            f'<img src="{url}" alt="流程图" style="max-width:100%;border:1px solid #e0e0e0;border-radius:6px;" />'
            f'<div style="font-size:11px;color:#999;margin-top:4px;">流程图</div></div>'
        )

    return _re.sub(r'```mermaid\s*\n(.*?)\n```', _render_mermaid, md_text, flags=_re.DOTALL)


def _md_to_html(md_text: str) -> str:
    """Markdown → HTML（含 echarts/mermaid 占位处理）。"""
    import re as _re
    # 替换 echarts → HTML 表格，mermaid → 图片
    cleaned = _render_echarts_blocks(md_text)
    cleaned = _render_mermaid_blocks(cleaned)
    try:
        import markdown as _md
        return _md.markdown(cleaned, extensions=['tables', 'fenced_code', 'codehilite', 'toc'])
    except ImportError:
        # 无 markdown 库时简单处理
        return f"<pre>{cleaned}</pre>"


def _build_docx_with_charts(md_text: str, title: str) -> bytes:
    """将 Markdown 报告转为 .docx，ECharts → 格式化表格，Mermaid → mermaid.ink SVG 图片。"""
    import re as _re
    import io, base64, zlib
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 标题
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 内联 Markdown → Word 格式化段落 ──
    def _add_rich_paragraph(text: str, style=None, indent=None, italic_all=False, color=None):
        """解析 **粗体** *斜体* `代码` [链接](url) <br> 并生成 Word 段落。"""
        # <br> → 换行
        text = _re.sub(r'<br\s*/?>', '\n', text, flags=_re.IGNORECASE)
        p = doc.add_paragraph(style=style)
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        if not text:
            return p

        # 正则匹配 inline 元素
        pattern = r'(\*\*(.+?)\*\*|'
        pattern += r'\*(.+?)\*|'
        pattern += r'`(.+?)`|'
        pattern += r'\[([^\]]+)\]\(([^)]+)\))'

        last = 0
        for m in _re.finditer(pattern, text):
            # 前面的普通文本
            if m.start() > last:
                plain = text[last:m.start()]
                if plain:
                    run = p.add_run(plain)
                    if italic_all:
                        run.italic = True
                    if color:
                        run.font.color.rgb = color

            # 粗体 **...**
            if m.group(2):
                run = p.add_run(m.group(2))
                run.bold = True
                if italic_all:
                    run.italic = True
                if color:
                    run.font.color.rgb = color
            # 斜体 *...*
            elif m.group(3):
                run = p.add_run(m.group(3))
                run.italic = True
                if color:
                    run.font.color.rgb = color
            # 代码 `...`
            elif m.group(4):
                run = p.add_run(m.group(4))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                if color:
                    run.font.color.rgb = color
            # 链接 [...](url)
            elif m.group(5):
                run = p.add_run(m.group(5))
                run.underline = True
                run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
                if italic_all:
                    run.italic = True

            last = m.end()

        # 尾部普通文本
        if last < len(text):
            plain = text[last:]
            if plain:
                run = p.add_run(plain)
                if italic_all:
                    run.italic = True
                if color:
                    run.font.color.rgb = color

        p.paragraph_format.space_after = Pt(6)
        return p

    # 按行解析 markdown，追踪状态
    lines = md_text.split('\n')
    in_code_block = False
    code_block_type = None
    code_block_lines = []
    in_table = False
    table_rows = []

    def _fill_cell_rich(cell, text: str, bold_all=False, font_size=Pt(10)):
        """填充单元格，解析内联 Markdown 格式和 <br> 换行。"""
        # 清除默认空段落
        for p in cell.paragraphs:
            p.clear()
        # <br> → 换行符
        text = _re.sub(r'<br\s*/?>', '\n', text, flags=_re.IGNORECASE)
        p = cell.paragraphs[0]
        # 匹配 **bold** *italic* `code`
        pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'
        last = 0
        for m in _re.finditer(pattern, text):
            if m.start() > last:
                run = p.add_run(text[last:m.start()])
                run.font.size = font_size
                if bold_all: run.bold = True
            if m.group(2):  # **bold**
                run = p.add_run(m.group(2))
                run.bold = True
                run.font.size = font_size
            elif m.group(3):  # *italic*
                run = p.add_run(m.group(3))
                run.italic = True
                run.font.size = font_size
            elif m.group(4):  # `code`
                run = p.add_run(m.group(4))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            last = m.end()
        if last < len(text):
            run = p.add_run(text[last:])
            run.font.size = font_size
            if bold_all: run.bold = True

    def _set_cell_margins(cell, top=60, bottom=60, left=80, right=80):
        """设置单元格内边距 + 垂直居中。"""
        from docx.oxml.ns import qn
        tc = cell._element.get_or_add_tcPr()
        # 垂直居中
        vAlign = tc.makeelement(qn('w:vAlign'), {qn('w:val'): 'center'})
        tc.append(vAlign)
        # 内边距
        tcMar = tc.makeelement(qn('w:tcMar'), {})
        for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            el = tcMar.makeelement(qn(f'w:{side}'), {qn('w:w'): str(val), qn('w:type'): 'dxa'})
            tcMar.append(el)
        tc.append(tcMar)

    def _flush_table():
        nonlocal in_table, table_rows
        if not table_rows or len(table_rows) < 2:
            table_rows = []
            in_table = False
            return
        header = [c.strip() for c in table_rows[0].split('|') if c.strip()]
        data_rows = []
        for tr in table_rows[1:]:
            cells = [c.strip() for c in tr.split('|') if c.strip()]
            if cells:
                data_rows.append(cells)
        if not header or not data_rows:
            table_rows = []
            in_table = False
            return
        tbl = doc.add_table(rows=1 + len(data_rows), cols=len(header), style='Table Grid')
        tbl.autofit = True
        from docx.oxml.ns import qn
        for j, h in enumerate(header):
            cell = tbl.rows[0].cells[j]
            _fill_cell_rich(cell, h, bold_all=True)
            _set_cell_margins(cell)
            shading = cell._element.get_or_add_tcPr()
            shd = shading.makeelement(qn('w:shd'), {qn('w:fill'): 'E8ECF1', qn('w:val'): 'clear'})
            shading.append(shd)
        for i, row in enumerate(data_rows):
            for j, val in enumerate(row):
                if j < len(header):
                    cell = tbl.rows[i + 1].cells[j]
                    _fill_cell_rich(cell, val)
                    _set_cell_margins(cell)
        doc.add_paragraph()
        table_rows = []
        in_table = False

    def _flush_code_block():
        nonlocal in_code_block, code_block_type, code_block_lines
        code = '\n'.join(code_block_lines)
        if code_block_type == 'echarts':
            _add_matplotlib_chart(doc, code)
        elif code_block_type == 'mermaid':
            _add_mermaid_image(doc, code)
        else:
            # 普通代码块
            p = doc.add_paragraph()
            run = p.add_run(code)
            run.font.size = Pt(9)
            run.font.name = 'Consolas'
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        code_block_lines = []
        code_block_type = None
        in_code_block = False

    for line in lines:
        # 代码块
        if line.strip().startswith('```'):
            if in_code_block:
                _flush_code_block()
            else:
                in_code_block = True
                code_block_type = line.strip()[3:].strip() or None
            continue

        if in_code_block:
            code_block_lines.append(line)
            continue

        # 表格
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if not in_table:
                _flush_table()  # flush any pending
                in_table = True
            # 跳过分隔行
            if not _re.match(r'^[\|\s\-:]+$', line.strip()):
                table_rows.append(line)
            continue
        elif in_table:
            _flush_table()

        # 空行
        if not line.strip():
            continue

        # 标题
        h_match = _re.match(r'^(#{1,6})\s+(.*)', line)
        if h_match:
            level = len(h_match.group(1))
            doc.add_heading(h_match.group(2), level=min(level, 3) + (0 if level <= 3 else 1))
            continue

        # 水平线
        if _re.match(r'^[-*_]{3,}$', line.strip()):
            doc.add_paragraph('─' * 50)
            continue

        # 列表（无序 + 有序）
        li_match = _re.match(r'^(\s*)[-*+]\s+(.*)', line)
        ol_match = _re.match(r'^(\s*)\d+[\.\)]\s+(.*)', line)
        if li_match:
            _add_rich_paragraph(li_match.group(2), style='List Bullet')
            continue
        if ol_match:
            _add_rich_paragraph(ol_match.group(2), style='List Number')
            continue

        # 引用
        bq_match = _re.match(r'^>\s?(.*)', line)
        if bq_match:
            _add_rich_paragraph(bq_match.group(1), indent=1, italic_all=True,
                               color=RGBColor(0x66, 0x66, 0x66))
            continue

        # 普通段落
        if line.strip():
            _add_rich_paragraph(line.strip())

    # Flush remaining
    if in_code_block:
        _flush_code_block()
    if in_table:
        _flush_table()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _add_matplotlib_chart(doc, echarts_code: str):
    """将 ECharts option 用 matplotlib 渲染为图表图片，嵌入 docx。"""
    import io
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from matplotlib.font_manager import FontProperties
    from docx.shared import Inches

    opt = _parse_echarts_option(echarts_code)
    if not opt:
        return

    # 中文字体（优先用项目内置 simhei.ttf）
    import os as _os
    zh_font = zh_font_title = None
    _proj_dir = _os.path.dirname(_os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))
    font_paths = [
        _os.path.join(_proj_dir, 'app', 'static', 'fonts', 'simhei.ttf'),  # 项目内置
        r'C:\Windows\Fonts\simhei.ttf',              # Windows
        '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',  # Linux fallback
    ]
    for fp in font_paths:
        if _os.path.isfile(fp):
            try:
                zh_font = FontProperties(fname=fp, size=10)
                zh_font_title = FontProperties(fname=fp, size=12)
                break
            except Exception:
                continue

    # 图表类型和标题
    series_list = opt.get("series", [])
    if isinstance(series_list, dict):
        series_list = [series_list]
    chart_type = series_list[0].get("type", "bar") if series_list else "bar"
    title_text = ""
    if "title" in opt:
        title_text = opt["title"].get("text", "") if isinstance(opt["title"], dict) else str(opt["title"])

    # 提取数据
    categories = []
    xaxis = opt.get("xAxis", {})
    if isinstance(xaxis, dict):
        categories = xaxis.get("data", [])
    elif isinstance(xaxis, list) and xaxis:
        categories = xaxis[0].get("data", []) if isinstance(xaxis[0], dict) else []

    # 辅助：提取数值
    def _to_num(v):
        if v is None: return 0
        if isinstance(v, (int, float)): return float(v)
        if isinstance(v, dict): return float(v.get("value", 0))
        try: return float(v)
        except: return 0

    # 辅助：提取标签
    def _to_label(v, default=""):
        if isinstance(v, dict): return str(v.get("name", default))
        return str(v)

    # 创建图表
    fig, ax = plt.subplots(figsize=(7, 3.5))
    colors = ['#5470c6', '#91cc75', '#fac858', '#ee6666', '#73c0de', '#3ba272', '#fc8452', '#9a60b4']

    if chart_type == "pie":
        # 饼图
        if series_list and series_list[0].get("data"):
            sdata = series_list[0]["data"]
            if isinstance(sdata[0], dict):
                values = [_to_num(d) for d in sdata]
                labels = [_to_label(d) for d in sdata]
            else:
                values = [_to_num(v) for v in sdata]
                labels = [str(c) for c in categories] if categories else [str(i+1) for i in range(len(values))]
            wedges, texts, autotexts = ax.pie(values, labels=labels, autopct='%1.1f%%',
                colors=colors[:len(values)], textprops={'fontsize': 8, 'fontproperties': zh_font} if zh_font else {'fontsize': 8})
            if autotexts:
                for at in autotexts:
                    at.set_fontsize(8)
        if title_text:
            ax.set_title(title_text, fontproperties=zh_font_title, fontsize=13, pad=12)

    elif chart_type == "line":
        # 折线图
        for idx, s in enumerate(series_list):
            sdata = s.get("data", [])
            vals = [_to_num(v) for v in sdata]
            if categories:
                ax.plot(range(len(vals)), vals, marker='o', color=colors[idx % len(colors)],
                       label=s.get("name", ""), linewidth=2, markersize=5)
                ax.set_xticks(range(len(categories)))
                ax.set_xticklabels(categories, fontproperties=zh_font, fontsize=8, rotation=30)
            else:
                ax.plot(vals, marker='o', color=colors[idx % len(colors)],
                       label=s.get("name", ""), linewidth=2, markersize=5)
        ax.legend(prop=zh_font, fontsize=8, loc='best')
        ax.grid(True, alpha=0.3)
        if title_text:
            ax.set_title(title_text, fontproperties=zh_font_title, fontsize=13, pad=12)

    else:
        # 柱状图 (默认)
        bar_width = 0.7 / len(series_list) if len(series_list) > 1 else 0.6
        x = range(len(categories)) if categories else range(max((len(s.get("data", [])) for s in series_list), default=0))
        for idx, s in enumerate(series_list):
            sdata = s.get("data", [])
            vals = [_to_num(v) for v in sdata]
            offset = (idx - (len(series_list) - 1) / 2) * bar_width if len(series_list) > 1 else 0
            positions = [i + offset for i in range(len(vals))]
            ax.bar(positions, vals, bar_width * 0.9, color=colors[idx % len(colors)],
                   label=s.get("name", ""), edgecolor='white', linewidth=0.5)
        if categories:
            ax.set_xticks(range(len(categories)))
            ax.set_xticklabels(categories, fontproperties=zh_font, fontsize=8, rotation=30)
        ax.legend(prop=zh_font, fontsize=8, loc='best') if len(series_list) > 1 else None
        ax.grid(True, axis='y', alpha=0.3)
        if title_text:
            ax.set_title(title_text, fontproperties=zh_font_title, fontsize=13, pad=12)

    plt.tight_layout()

    # 保存为 PNG 字节
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white', edgecolor='none')
    plt.close(fig)
    buf.seek(0)

    # 嵌入 docx
    doc.add_paragraph()  # spacing
    doc.add_picture(buf, width=Inches(5.5))
    last_para = doc.paragraphs[-1]
    last_para.alignment = 1  # center
    doc.add_paragraph()
    buf.close()


def _add_mermaid_image(doc, mermaid_code: str):
    """将 Mermaid 流程图嵌入 docx。优先级: mmdc 本地 → mermaid.ink → 美化代码块。

    内网部署无外网时依赖 mmdc CLI (@mermaid-js/mermaid-cli)。
    安装: npm install -g @mermaid-js/mermaid-cli
    """
    import base64, io, shutil, subprocess, tempfile

    from docx.shared import Inches, Pt, RGBColor

    img_bytes = None

    # ── 1) mmdc 本地渲染（无网络依赖）──
    mmdc = shutil.which('mmdc')
    if mmdc:
        try:
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                tmp_path = tmp.name
            # 写 mermaid 代码到临时 .mmd 文件（mmdc 对 stdin 支持有差异，文件更稳定）
            mmd_path = tmp_path + '.mmd'
            with open(mmd_path, 'w', encoding='utf-8') as f:
                f.write(mermaid_code)
            result = subprocess.run(
                [mmdc, '-i', mmd_path, '-o', tmp_path, '-s', '2', '-b', 'white'],
                capture_output=True, text=True, timeout=30,
            )
            if result.returncode == 0:
                with open(tmp_path, 'rb') as f:
                    img_bytes = f.read()
            else:
                # mmdc 失败时 stderr 常有有用信息
                from loguru import logger
                logger.warning(f"[mmdc] 渲染失败:\n{result.stderr[:500]}")
        except Exception as exc:
            from loguru import logger
            logger.warning(f"[mmdc] 调用异常: {exc}")
        finally:
            for p in (tmp_path, mmd_path):
                try:
                    import os
                    os.unlink(p)
                except Exception:
                    pass

    # ── 2) mermaid.ink 在线渲染（兜底）──
    if img_bytes is None:
        try:
            import requests as _req
            encoded = base64.urlsafe_b64encode(mermaid_code.encode("utf-8")).decode("ascii").rstrip("=")
            url = f"https://mermaid.ink/img/{encoded}"
            resp = _req.get(url, timeout=15)
            if resp.status_code == 200:
                img_bytes = resp.content
        except Exception:
            pass

    # ── 3) 有图片 → 嵌入 docx ──
    if img_bytes:
        try:
            img_stream = io.BytesIO(img_bytes)
            doc.add_picture(img_stream, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = 1  # center
            return
        except Exception:
            pass

    # ── 4) 全部失败 → 美化代码块嵌入 ──
    p = doc.add_paragraph()
    run = p.add_run('▶ 流程图')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1a, 0x5c, 0x8a)
    p.paragraph_format.space_after = Pt(4)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(mermaid_code)
    run2.font.size = Pt(9)
    run2.font.name = 'Consolas'
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p2.paragraph_format.space_after = Pt(8)


def _html_to_docx(html_doc: str, title: str) -> bytes:
    """HTML → .docx 字节（使用 python-docx 构建基础文档）。"""
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re as _re

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 标题
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 简单 HTML → docx 段落转换
    # 按块级元素分割
    blocks = _re.split(r'</?(?:h[1-6]|p|ul|ol|li|table|tr|pre|blockquote|hr|div)[^>]*>', html_doc)
    # 更简单的办法：去掉 HTML 标签，保留文本结构
    # 按标题标记分段
    lines = html_doc.split('\n')
    in_table = False
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # 标题
        for level in range(1, 7):
            htag = f'<h{level}>'
            if stripped.startswith(htag):
                text = _re.sub(r'<[^>]+>', '', stripped)
                doc.add_heading(text, level=level)
                break
        else:
            # 表格
            if '<table>' in stripped or '<tr>' in stripped:
                in_table = True
                continue
            if '</table>' in stripped:
                in_table = False
                continue
            if in_table:
                continue

            # 去掉所有 HTML 标签
            text = _re.sub(r'<[^>]+>', '', stripped)
            if text:
                para = doc.add_paragraph(text)
                style = para.paragraph_format
                style.space_after = Pt(6)

    # 保存到内存
    import io
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _url_quote(s: str) -> str:
    """URL 编码文件名（UTF-8）。"""
    from urllib.parse import quote
    return quote(s, safe='')

@router.get("/processed", summary="获取已处理审批列表")
async def get_processed_confirmations(
    page: int = 1,
    page_size: int = 20,
    db: AsyncSession = Depends(get_db),
):
    """获取已处理的审批列表（通过/拒绝），分页。"""
    from app.repositories.message_repository import MessageRepository
    repo = MessageRepository(db)
    offset = max(0, (page - 1)) * max(1, page_size)
    items = await repo.get_processed_confirmations(
        limit=max(1, page_size), offset=offset,
    )
    total_count = await repo.count_processed_confirmations()
    result = []
    for msg in items:
        content_data = {}
        try:
            content_data = json.loads(msg.content) if msg.content else {}
        except Exception:
            content_data = {}
        result.append({
            "id": msg.id,
            "conversation_id": msg.conversation_id,
            "conversation_title": content_data.get("conversation_title", ""),
            "message_type": msg.message_type or "",
            "action_label": content_data.get("action_label", ""),
            "concept_label": content_data.get("concept_label", ""),
            "tool": content_data.get("tool", ""),
            "params": content_data.get("params", {}),
            "risk": content_data.get("risk", "write"),
            "user_id": content_data.get("user_id", ""),
            "param_schema": content_data.get("param_schema", []),
            "error_detail": content_data.get("error_detail", ""),
            "decision_pack": content_data.get("decision_pack", {}),
            "verify_detail": content_data.get("verify_detail", []),
            "verify_summary": content_data.get("verify_summary", ""),
            "submitter_id": content_data.get("submitter_id", ""),
            "chain_id": content_data.get("chain_id", ""),
            "message_id": content_data.get("message_id", ""),
            "assigned_to": msg.assigned_to or "",
            "status": msg.status,
            "reviewed_by": msg.reviewed_by or "",
            "reviewed_at": msg.reviewed_at or "",
            "created_at": str(msg.created_at) if msg.created_at else "",
        })
    # 审批条目（confirm）content 无 conversation_title：从会话表回退查，保证「📎 原对话」入口可用
    _missing = [it for it in result if not it.get("conversation_title") and it.get("conversation_id")]
    if _missing:
        from app.repositories.conversation_repository import ConversationRepository
        _cr = ConversationRepository(db)
        for _it in _missing:
            try:
                _cv = await _cr.get_by_id(_it["conversation_id"])
                if _cv and _cv.title:
                    _it["conversation_title"] = _cv.title
            except Exception:
                pass
    return {
        "processed": result, "total": total_count,
        "page": page, "page_size": page_size,
    }


class BatchRequest(BaseModel):
    message_ids: list[str]
    user_id: str = ""
    comment: str = ""


@router.post("/batch-approve", summary="批量通过审批")
async def batch_approve_confirmations(
    body: BatchRequest,
    db: AsyncSession = Depends(get_db),
):
    """批量通过审批。每条用独立 session 避免 commit 后事务失效。"""
    from app.db import _async_session as _session_factory
    success, failed = 0, []
    for mid in body.message_ids:
        try:
            async with _session_factory() as session:
                req = ApprovalRequest(user_id=body.user_id, comment=body.comment)
                await approve_confirmation(mid, req, session)
            success += 1
        except HTTPException:
            failed.append({"id": mid, "error": "消息不存在或已处理"})
        except Exception as e:
            failed.append({"id": mid, "error": str(e)})
    return {"success": success, "failed": failed}


@router.post("/batch-reject", summary="批量拒绝审批")
async def batch_reject_confirmations(
    body: BatchRequest,
    db: AsyncSession = Depends(get_db),
):
    """批量拒绝审批。每条用独立 session 避免 commit 后事务失效。"""
    from app.db import _async_session as _session_factory
    success, failed = 0, []
    for mid in body.message_ids:
        try:
            async with _session_factory() as session:
                req = ApprovalRequest(user_id=body.user_id, comment=body.comment)
                await reject_confirmation(mid, req, session)
            success += 1
        except HTTPException:
            failed.append({"id": mid, "error": "消息不存在或已处理"})
        except Exception as e:
            failed.append({"id": mid, "error": str(e)})
    return {"success": success, "failed": failed}


class ExecutePlanRequest(BaseModel):
    chain_id: str
    params: dict = {}
    conversation_id: str = ""
    message_id: str = ""  # 精确指定保存目标消息（前端方案所在消息）


@router.post("/execute-plan", summary="执行变更方案")
async def execute_change_plan(
    body: ExecutePlanRequest,
    http_request: Request,
):
    """用户选择变更方案后，加载对应执行链并逐步执行。"""
    from app.core.chain_engine import chain_engine

    chain_result = {"ok": 0, "total": 0, "errors": []}

    async def event_generator():
        nonlocal chain_result
        async for chunk_type, chunk_content in chain_engine.execute(
            message="执行变更方案",
            chain_id=body.chain_id,
            session_id=body.conversation_id,
            params=body.params,
        ):
            yield f"data: {json.dumps({'type': chunk_type, 'content': chunk_content}, ensure_ascii=False)}\n\n"
            if chunk_type == 'chain_done':
                try:
                    cd = json.loads(chunk_content) if isinstance(chunk_content, str) else chunk_content
                    chain_result = {
                        "ok": cd.get("steps_completed", 0),
                        "total": cd.get("total_steps", 0),
                        "verified": cd.get("verified"),
                        "verify_summary": cd.get("verify_summary", ""),
                        "verify_detail": cd.get("verify_detail", []) or [],
                        "review_required": cd.get("review_required", False),
                        "rolled_back": cd.get("rolled_back", False),
                    }
                except Exception:
                    pass
            elif chunk_type == 'chain_step':
                try:
                    cs = json.loads(chunk_content) if isinstance(chunk_content, str) else chunk_content
                    if cs.get("status") == "error":
                        chain_result["errors"].append(cs.get("description", ""))
                except Exception:
                    pass
            elif chunk_type == 'verify_result':
                # verify 阶段结果：verified / detail / review_required / rolled_back
                try:
                    vr = json.loads(chunk_content) if isinstance(chunk_content, str) else chunk_content
                    chain_result["verified"] = vr.get("verified")
                    chain_result["verify_summary"] = vr.get("summary", "")
                    chain_result["verify_detail"] = vr.get("detail", []) or []
                    chain_result["review_required"] = vr.get("review_required", False)
                    chain_result["rolled_back"] = vr.get("rolled_back", False)
                except Exception:
                    pass

        # 写入执行结果到对话：更新最新 assistant 消息的 metadata
        print(f"[ExecutePlan DEBUG] conv_id={body.conversation_id}, chain_id={body.chain_id}, result={chain_result}", flush=True)
        if not body.conversation_id:
            print("[ExecutePlan DEBUG] SKIP: no conversation_id", flush=True)
        else:
            try:
                from app.db import _async_session as _sf
                from app.repositories.message_repository import MessageRepository
                from app.models.message import MessageRole
                async with _sf() as session:
                    repo = MessageRepository(session)
                    target_msg = None
                    if body.message_id:  # 前端精确指定保存目标消息
                        target_msg = await repo.get_by_id(body.message_id)
                    if not target_msg:  # 回退：最新 assistant
                        msgs = await repo.get_by_conversation(body.conversation_id, limit=20, offset=0)
                        for m in reversed(msgs):
                            if m.role == MessageRole.ASSISTANT:
                                target_msg = m
                                break
                    if target_msg:
                        print(f"[ExecutePlan DEBUG] Found target_msg={target_msg.id}, saving...", flush=True)
                        meta = target_msg.metadata_dict or {}
                        exec_results = meta.get('plan_exec_results', {})
                        err_text = f"，{len(chain_result['errors'])}步失败" if chain_result["errors"] else ""
                        _v_verified = chain_result.get('verified')
                        _v_status = (
                            'needs_review' if _v_verified is False
                            else ('failed' if chain_result['errors'] else 'ok')
                        )
                        exec_results[body.chain_id] = {
                            'status': _v_status,
                            'ok': chain_result['ok'],
                            'total': chain_result['total'],
                            'summary': f"{chain_result['ok']}/{chain_result['total']} 成功{err_text}",
                            'verified': _v_verified,
                            'verify_summary': chain_result.get('verify_summary', ''),
                        }
                        meta['plan_exec_results'] = exec_results
                        target_msg.metadata_dict = meta
                        await session.commit()
                        log.info(f"[ExecutePlan] 执行结果已写入消息 {target_msg.id}: {body.chain_id} → {exec_results[body.chain_id]['status']}")
                    else:
                        print(f"[ExecutePlan DEBUG] No assistant msg found in conv={body.conversation_id}", flush=True)
            except Exception as e:
                print(f"[ExecutePlan DEBUG] ERROR: {e}", flush=True)
                import traceback; traceback.print_exc()

        # ── 变更类写操作验证失败 → 创建责任分离复核条目（他人按角色复核）──
        # 与自动回滚互斥：已自动回滚（rolled_back=True）不创建；分析类（review_required=False）就地标记不进队列
        if (chain_result.get("verified") is False and chain_result.get("review_required")
                and not chain_result.get("rolled_back")):
            try:
                from app.models.message import MessageType, ConfirmStatus, MessageRole
                from app.services.event_bus import event_bus
                from app.core.chain_engine import _CHAINS
                from app.services.action_executor import action_executor
                # 复核角色：链中写 action 的 review_roles 合并（空=不限制）
                review_roles = []
                try:
                    action_executor._ensure_loaded()
                    cfg = _CHAINS.get(body.chain_id) or {}
                    for rs in (cfg.get("reasoning_steps") or []):
                        an = rs.get("action_name", "") if isinstance(rs, dict) else ""
                        if not an:
                            continue
                        sig = action_executor._sigs.get(an) or {}
                        for r in (sig.get("review_roles") or []):
                            if r and r not in review_roles:
                                review_roles.append(r)
                except Exception:
                    pass
                _submitter = ""
                try:
                    _submitter = get_current_user_id(http_request)
                except Exception:
                    _submitter = ""
                _chain_name = ""
                try:
                    _chain_name = (_CHAINS.get(body.chain_id) or {}).get("name", "")
                except Exception:
                    pass
                # 定位原对话：body.conversation_id 优先，为空则从方案所在消息反查（防会话 ID 丢失）
                _conv_id = body.conversation_id or ""
                _conv_title = ""
                try:
                    async with _sf() as session:
                        repo0 = MessageRepository(session)
                        if not _conv_id and body.message_id:
                            _m0 = await repo0.get_by_id(body.message_id)
                            if _m0:
                                _conv_id = _m0.conversation_id or ""
                        if _conv_id:
                            from app.repositories.conversation_repository import ConversationRepository
                            _cv = await ConversationRepository(session).get_by_id(_conv_id)
                            _conv_title = _cv.title if _cv else ""
                except Exception:
                    pass
                async with _sf() as session:
                    repo = MessageRepository(session)
                    review_msg = await repo.create(
                        conversation_id=_conv_id,
                        role=MessageRole.SYSTEM,
                        content=json.dumps({
                            "review_type": "verify_failed",
                            "chain_id": body.chain_id,
                            "chain_name": _chain_name,
                            "action_label": _chain_name,
                            "concept_label": "",
                            "verify_summary": chain_result.get("verify_summary", ""),
                            "verify_detail": chain_result.get("verify_detail") or [],
                            "submitter_id": _submitter,
                            "conversation_id": _conv_id,
                            "conversation_title": _conv_title,
                            "message_id": body.message_id or "",
                            "message": f"变更方案 {_chain_name or body.chain_id} 执行后验证未通过，待复核",
                        }, ensure_ascii=False),
                        message_type=MessageType.REVIEW.value,
                        status=ConfirmStatus.PENDING.value,
                        assigned_to=",".join(review_roles),
                    )
                    await session.commit()
                    review_id = review_msg.id
                await event_bus.publish("pending_updated", {"source": "execute_plan"})
                yield f"data: {json.dumps({'type': 'review_created', 'content': {'message_id': review_id, 'verify_summary': chain_result.get('verify_summary', '')}}, ensure_ascii=False)}\n\n"
                log.info(f"[ExecutePlan] 验证失败已创建复核条目 {review_id}，复核角色={review_roles}")
            except Exception as e:
                log.error(f"[ExecutePlan] 创建复核条目失败: {e}")
        yield "data: [DONE]\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


class SavePlanResultRequest(BaseModel):
    conversation_id: str
    chain_id: str
    status: str  # 'ok' | 'failed' | 'needs_review'
    ok: int = 0
    total: int = 0
    summary: str = ""
    verified: Optional[bool] = None
    verify_summary: str = ""
    verify_detail: list = []  # [{property, expected, actual, match}]
    message_id: str = ""  # 精确指定保存目标消息（前端方案所在消息）


@router.post("/save-plan", summary="保存方案执行结果到消息 metadata")
async def save_plan_result(body: SavePlanResultRequest):
    """前端执行完成后调用，将结果写入最新 assistant 消息的 metadata。"""
    if not body.conversation_id:
        return {"ok": False, "error": "conversation_id 为空"}
    try:
        from app.db import _async_session as _sf
        from app.repositories.message_repository import MessageRepository
        from app.models.message import MessageRole
        async with _sf() as session:
            repo = MessageRepository(session)
            target_msg = None
            if body.message_id:  # 前端精确指定保存目标消息
                target_msg = await repo.get_by_id(body.message_id)
            if not target_msg:  # 回退：最新 assistant
                msgs = await repo.get_by_conversation(body.conversation_id, limit=20, offset=0)
                for m in reversed(msgs):
                    if m.role == MessageRole.ASSISTANT:
                        target_msg = m
                        break
            if target_msg:
                meta = target_msg.metadata_dict or {}
                exec_results = meta.get('plan_exec_results', {})
                exec_results[body.chain_id] = {
                    'status': body.status,
                    'ok': body.ok,
                    'total': body.total,
                    'summary': body.summary,
                    'verified': body.verified,
                    'verify_summary': body.verify_summary,
                    'verify_detail': body.verify_detail or [],
                }
                meta['plan_exec_results'] = exec_results
                target_msg.metadata_dict = meta
                await session.commit()
                log.info(f"[SavePlanResult] 已保存: msg={target_msg.id} chain={body.chain_id} status={body.status}")
                # 审计账本（append-only）：记录执行链结果，供追溯
                try:
                    from app.agents.guardrails import AuditLogger
                    AuditLogger.log(
                        tool_name="chain_execute",
                        action_name=body.chain_id,
                        risk="medium",
                        agent="chain_engine",
                        params={"conversation_id": body.conversation_id, "message_id": body.message_id or ""},
                        result_preview=body.verify_summary or body.summary,
                        success=body.verified is not False,
                        session_id=body.conversation_id,
                    )
                except Exception:
                    pass
                return {"ok": True}
            else:
                return {"ok": False, "error": "未找到 assistant 消息"}
    except Exception as e:
        log.error(f"[SavePlanResult] 失败: {e}")
        return {"ok": False, "error": str(e)}


@router.post("/restore-entity", summary="轨迹回滚：恢复被删实体 / 删除新建实体")
async def restore_entity(request: RestoreEntityRequest):
    """轨迹回滚：按执行轨迹恢复。

    - 工具名以 _delete 结尾 → 回滚 = 恢复被删实体（用 before_snapshot 重新创建）
    - 工具名以 _create/_write 结尾 → 回滚 = 删除新建实体（用 created_entity_id/主键）
    """
    from app.services.action_executor import action_executor
    from app.services.data_backend import data_backend
    from app.services.ontology_service import ontology_service

    action_executor._ensure_loaded()
    sig = action_executor._sigs.get(request.tool) or {}
    concept = sig.get("conceptName", "") or request.tool.replace("_delete", "").replace("_create", "").replace("_update", "")
    if not concept:
        return {"ok": False, "error": f"无法从工具 {request.tool} 确定概念"}
    is_create = sig.get("actionName") == "create" or request.tool.endswith("_create")

    # 主键名
    pk_name = "code"
    try:
        cdef = ontology_service.get_concept(concept)
        if cdef:
            for p in cdef.get("properties", []):
                if p.get("isPrimary"):
                    pk_name = p.get("name", "code")
                    break
    except Exception:
        pass

    if is_create:
        # 回滚 = 删除新建实体
        deleted = 0
        for record in request.records:
            pk_val = record.get(pk_name)
            if not pk_val and request.created_entity_id:
                pk_val = request.created_entity_id
            if pk_val and await data_backend.delete(concept, pk_name, str(pk_val)):
                deleted += 1
        return {"ok": True, "concept": concept, "deleted": deleted, "total": len(request.records)}
    else:
        # 回滚 = 恢复被删实体
        restored, failed = 0, 0
        for record in request.records:
            try:
                result = await data_backend.create(concept, record)
                if result and not result.get("error"):
                    restored += 1
                else:
                    failed += 1
            except Exception:
                failed += 1
        return {"ok": True, "concept": concept, "restored": restored, "failed": failed, "total": len(request.records)}


@router.get("/audit/logs", summary="获取执行链审计日志")
async def get_audit_logs(limit: int = 100, keyword: str = ""):
    """读取 append-only 审计日志（logs/audit.log），返回最近记录。

    审计日志由 AuditLogger 追加写入（含 chain_execute 执行链记录：
    chain_id、verified、verify_summary、时间、会话）。
    """
    from app.agents.settings import AUDIT_CONFIG
    log_file = AUDIT_CONFIG.get("log_file", "logs/audit.log")
    logs = []
    try:
        with open(log_file, "r", encoding="utf-8") as f:
            lines = f.readlines()
        # 倒序：最新记录在前（文件尾部是最新，反转后第一条即最新）
        for line in lines[-limit:][::-1]:
            try:
                entry = json.loads(line)
                if keyword and keyword not in json.dumps(entry, ensure_ascii=False):
                    continue
                logs.append(entry)
            except Exception:
                continue
    except Exception:
        pass
    return {"logs": logs, "total": len(logs)}


@router.get("/prompt-logs", summary="获取提示词日志")
async def get_prompt_logs(
    page: int = 1, page_size: int = 20,
    keyword: str = "",
    db: AsyncSession = Depends(get_db),
):
    """获取 LLM 提示词记录（分页+搜索）。"""
    from app.repositories.message_repository import MessageRepository
    repo = MessageRepository(db)
    offset = max(0, (page - 1)) * max(1, page_size)
    all_msgs = await repo.get_latest_with_metadata(limit=500, offset=0)
    logs = []
    for msg in all_msgs:
        meta = msg.metadata_dict or {}
        pi = meta.get("prompt_info")
        if not pi:
            continue
        user_msg = pi.get("user_message", "")
        sp = pi.get("system_prompt", "")
        if keyword and keyword not in user_msg and keyword not in sp:
            continue
        logs.append({
            "id": msg.id,
            "conversation_id": msg.conversation_id,
            "created_at": str(msg.created_at) if msg.created_at else "",
            "model": pi.get("model", ""),
            "system_prompt": sp,
            "system_prompt_len": len(sp),
            "user_message": user_msg,
            "enable_thinking": pi.get("enable_thinking", False),
            "web_search": pi.get("web_search", False),
            "input_tokens": pi.get("input_tokens", 0),
            "output_tokens": pi.get("output_tokens", 0),
        })
    total = len(logs)
    paged = logs[offset:offset + max(1, page_size)]
    return {"logs": paged, "total": total, "page": page, "page_size": page_size}


@router.delete("/batch", summary="批量删除消息")
async def batch_delete_messages(
    body: BatchRequest,
):
    """批量删除消息记录。"""
    from app.db import _async_session as _session_factory
    from app.repositories.message_repository import MessageRepository
    async with _session_factory() as session:
        repo = MessageRepository(session)
        deleted = await repo.bulk_delete(body.message_ids)
    return {"deleted": deleted}


@router.post("/{message_id}/approve")
async def approve_confirmation(
    message_id: str,
    body: ApprovalRequest,
    db: AsyncSession = Depends(get_db),
):
    """通过审批并执行原始动作。"""
    from app.repositories.message_repository import MessageRepository
    from app.models.message import MessageType, ConfirmStatus, MessageRole

    repo = MessageRepository(db)
    pending_msg = await repo.get_by_id(message_id)
    if not pending_msg:
        raise HTTPException(status_code=404, detail=f"消息不存在: {message_id}")
    if pending_msg.status != ConfirmStatus.PENDING.value:
        raise HTTPException(status_code=400, detail="该消息已处理")
    if pending_msg.message_type == MessageType.REVIEW.value:
        raise HTTPException(status_code=400, detail="复核条目请使用复核动作（接受/回滚）")

    # 提取原始动作参数
    content_data = {}
    try:
        content_data = json.loads(pending_msg.content) if pending_msg.content else {}
    except Exception:
        pass
    tool_name = content_data.get("tool", "")
    params = content_data.get("params", {})
    original_user_id = content_data.get("user_id", "")
    conversation_id = pending_msg.conversation_id

    # 标记为已审批
    updated = await repo.resolve_confirmation(message_id, approved=True, reviewed_by=body.user_id)
    reviewer = body.user_id or "审批人"
    log.info(f"[审批] message_id={message_id} 已通过")

    # 执行原始动作 + 更新思考链（异常工单跳过执行）
    is_exception = content_data.get("risk") == "exception"
    exec_result = {"success": False, "message": "未执行", "rowCount": 0}
    action_label = content_data.get("action_label", tool_name)
    if not is_exception:
        try:
            if tool_name:
                from app.services.action_executor import action_executor
                exec_result = await action_executor.execute_structured_async(
                    tool_name, {**params, '_skip_approval': True}, user_id=original_user_id or body.user_id,
                )
                log.info(f"[审批] 动作 {tool_name} 执行完成: rowCount={exec_result.get('rowCount', 0)}")
        except Exception as e:
            log.error(f"[审批] 动作执行失败: {e}")
            exec_result = {"success": False, "message": str(e), "rowCount": 0}

    try:
        # 用 param_schema 翻译字段名为中文标签
        schema_map = {}
        try:
            cdata = json.loads(pending_msg.content) if isinstance(pending_msg.content, str) else (pending_msg.content or {})
            for ps in cdata.get("param_schema", []):
                schema_map[ps.get("name", "")] = ps.get("label", ps.get("name", ""))
        except Exception:
            pass
        labeled_params = {schema_map.get(k, k): v for k, v in params.items() if v}

        params_summary = ", ".join(f"{k}={v}" for k, v in labeled_params.items())
        await _append_exec_step(db, pending_msg, f"审批通过 ({reviewer})，已执行",
            f"审批人: {reviewer} | 操作: {action_label}" +
            (f" | {params_summary}" if params_summary else "") +
            (f" | 影响 {exec_result.get('rowCount', 0)} 行" if exec_result.get('rowCount', 0) > 0 else " | 完成")
        )
    except Exception as e:
        log.error(f"[审批] 更新思考链失败: {e}")

    # 执行结果写入对话（含审批人、参数摘要、执行结果）
    try:
        # 构建完整审批通知（对齐复核通知格式：动作/审批人/提交人/参数/结果/锚点）
        _conc = content_data.get("concept_label", "")
        _sub = content_data.get("user_id", "")
        result_parts = [
            f"✅ 审批通过：**{action_label}**" + (f"（{_conc}）" if _conc else "") + f"，由 **{reviewer}** 审批"
        ]
        if _sub:
            result_parts.append(f"提交人: {_sub}")
        # 参数摘要（中文标签）
        try:
            _sm = {}
            _cd = json.loads(pending_msg.content) if isinstance(pending_msg.content, str) else (pending_msg.content or {})
            for ps in _cd.get("param_schema", []):
                _sm[ps.get("name", "")] = ps.get("label", ps.get("name", ""))
            _lp = {_sm.get(k, k): v for k, v in params.items() if v}
            if _lp:
                result_parts.append("**参数**: " + ", ".join(f"{k}={v}" for k, v in _lp.items()))
        except Exception:
            pass
        # 执行结果
        if is_exception:
            result_parts.append("异常工单已处理，未执行数据变更")
        else:
            rc = exec_result.get("rowCount", 0)
            result_parts.append(f"**执行结果**: {'✅ 成功' if exec_result.get('success', True) and not exec_result.get('error') else '❌ 失败'}" + (f"，影响 {rc} 行" if rc else ""))
            if exec_result.get("error"):
                result_parts.append(f"> {exec_result['error']}")
            elif exec_result.get("result"):
                result_parts.append(f"\n{exec_result['result']}")
        # 审批通知带「打开原对话」锚点（点通知右侧抽屉看上下文）
        try:
            if conversation_id:
                result_parts.append(f"\n📎 [打开原对话](conv://{conversation_id})")
        except Exception:
            pass
        await repo.create(
            conversation_id=conversation_id, role=MessageRole.ASSISTANT,
            content="\n".join(result_parts), message_type=MessageType.INFO.value,
        )
    except Exception as e:
        log.error(f"[审批] 写入对话失败: {e}")

    # 处理推理链确认
    inferences = exec_result.get("inferences", []) or []
    if exec_result.get("needs_inference_confirmation") and inferences:
        try:
            for inf in inferences:
                await repo.create(
                    conversation_id=conversation_id, role=MessageRole.SYSTEM,
                    content=json.dumps({
                        "tool": inf.get("target_action", tool_name),
                        "action_label": inf.get("rule_label", ""),
                        "concept_label": inf.get("target_concept", ""),
                        "params": inf.get("target_params", {}),
                        "risk": "inference", "user_id": body.user_id,
                        "message": f"推理链: {inf.get('description', '')}",
                    }, ensure_ascii=False),
                    message_type=MessageType.CONFIRM.value,
                    status=ConfirmStatus.PENDING.value,
                    assigned_to=content_data.get("assigned_to", ""),
                )
        except Exception as e:
            log.error(f"[审批] 写入推理链确认失败: {e}")

    # 广播审批完成事件
    from app.services.event_bus import event_bus
    await event_bus.publish("approval_done", {
        "conversation_id": conversation_id,
        "message_id": message_id,
        "action": action_label,
        "reviewer": reviewer,
        "submitter": original_user_id or "",
        "approved": True,
    })

    return {
        "success": True,
        "message_id": message_id,
        "status": updated.status,
        "exec_result": {"rowCount": exec_result.get("rowCount", 0), "message": exec_result.get("message", "")},
    }


async def _append_exec_step(db, pending_msg, label: str, detail: str = None, status: str = "done"):
    """往待审批消息关联的 AI 消息的思考链中追加执行步骤。"""
    try:
        from app.models.message import Message, MessageRole
        from sqlalchemy import select
        import json as _json
        q = select(Message).where(
            Message.conversation_id == pending_msg.conversation_id,
            Message.role == MessageRole.ASSISTANT,
        ).order_by(Message.created_at.desc()).limit(1)
        r = await db.execute(q)
        ai_msg = r.scalar_one_or_none()
        if ai_msg and ai_msg.extra_data:
            meta = _json.loads(ai_msg.extra_data) if isinstance(ai_msg.extra_data, str) else ai_msg.extra_data
            steps = meta.get("execution_steps", [])
            steps.append({
                "key": "approval_executed",
                "label": label,
                "status": status,
                "detail": detail if isinstance(detail, str) else _json.dumps(detail, ensure_ascii=False) if detail else "",
            })
            meta["execution_steps"] = steps
            ai_msg.extra_data = _json.dumps(meta, ensure_ascii=False)
            await db.commit()
            log.info(f"[审批] 思考链已更新: {label}")
    except Exception as e:
        log.warning(f"[审批] 更新思考链失败: {e}")


@router.post("/{message_id}/reject")
async def reject_confirmation(
    message_id: str,
    body: ApprovalRequest,
    db: AsyncSession = Depends(get_db),
):
    """拒绝审批。"""
    from app.repositories.message_repository import MessageRepository
    from app.models.message import MessageType, MessageRole, ConfirmStatus

    repo = MessageRepository(db)
    pending_msg = await repo.get_by_id(message_id)
    if not pending_msg:
        raise HTTPException(status_code=404, detail=f"消息不存在: {message_id}")
    if pending_msg.status != ConfirmStatus.PENDING.value:
        raise HTTPException(status_code=400, detail="该消息已处理")
    if pending_msg.message_type == MessageType.REVIEW.value:
        raise HTTPException(status_code=400, detail="复核条目请使用复核动作（接受/回滚）")

    # 提取动作信息
    content_data = {}
    try:
        content_data = json.loads(pending_msg.content) if pending_msg.content else {}
    except Exception:
        pass
    action_label = content_data.get("action_label", "")
    reason = body.comment or ""

    updated = await repo.resolve_confirmation(message_id, approved=False, reviewed_by=body.user_id)
    if not updated:
        raise HTTPException(status_code=404, detail=f"消息不存在: {message_id}")

    reviewer = body.user_id or "审批人"
    detail = f"审批人: {reviewer} | 操作: {action_label}"
    if reason:
        detail += f" | 原因: {reason}"
    try:
        await _append_exec_step(db, pending_msg, f"审批拒绝 ({reviewer})", detail, status="error")
    except Exception as e:
        log.error(f"[审批] 更新思考链失败: {e}")

    # 拒绝结果写入对话（完整信息，对齐复核通知格式）
    try:
        _conc = content_data.get("concept_label", "")
        _sub = content_data.get("user_id", "")
        reject_msg = (f"❌ 审批拒绝：**{action_label}**" + (f"（{_conc}）" if _conc else "") + f"，由 **{reviewer}** 拒绝")
        if _sub:
            reject_msg += f"\n提交人: {_sub}"
        if reason:
            reject_msg += f"\n原因: {reason}"
        # 审批通知带「打开原对话」锚点（点通知右侧抽屉看上下文）
        try:
            if pending_msg.conversation_id:
                reject_msg += f"\n📎 [打开原对话](conv://{pending_msg.conversation_id})"
        except Exception:
            pass
        await repo.create(
            conversation_id=pending_msg.conversation_id, role=MessageRole.ASSISTANT,
            content=reject_msg, message_type=MessageType.INFO.value,
        )
    except Exception as e:
        log.error(f"[审批] 写入对话失败: {e}")

    try:
        from app.services.event_bus import event_bus
        await event_bus.publish("approval_done", {
            "conversation_id": pending_msg.conversation_id,
            "message_id": message_id,
            "action": action_label,
            "reviewer": reviewer,
            "approved": False,
        })
    except Exception:
        pass

    return {"success": True, "message_id": message_id, "status": updated.status}


# ── 责任分离复核（验证失败后由复核角色接受或回滚）────────────────────

def _load_review_entry(pending_msg):
    """解析复核条目 content JSON + 责任分离校验（执行人不能复核自己的变更）。"""
    content_data = {}
    try:
        content_data = json.loads(pending_msg.content) if pending_msg.content else {}
    except Exception:
        content_data = {"raw": pending_msg.content}
    return content_data


@router.post("/{message_id}/review-accept")
async def review_accept(
    message_id: str,
    body: ApprovalRequest,
    db: AsyncSession = Depends(get_db),
):
    """复核人确认接受：验证失败结果可接受，标记复核通过（不重新执行）。"""
    from app.repositories.message_repository import MessageRepository
    from app.models.message import MessageType, ConfirmStatus, MessageRole

    repo = MessageRepository(db)
    pending_msg = await repo.get_by_id(message_id)
    if not pending_msg:
        raise HTTPException(status_code=404, detail=f"消息不存在: {message_id}")
    if pending_msg.message_type != MessageType.REVIEW.value:
        raise HTTPException(status_code=400, detail="非复核条目")
    if pending_msg.status != ConfirmStatus.PENDING.value:
        raise HTTPException(status_code=400, detail="该复核已处理")

    content_data = _load_review_entry(pending_msg)
    submitter = content_data.get("submitter_id", "")
    if submitter and body.user_id and submitter == body.user_id:
        raise HTTPException(status_code=400, detail="执行人不能复核自己的变更（责任分离）")

    reviewer = body.user_id or "复核人"
    updated = await repo.resolve_confirmation(message_id, approved=True, reviewed_by=reviewer)
    if not updated:
        raise HTTPException(status_code=404, detail=f"消息不存在: {message_id}")
    log.info(f"[复核] message_id={message_id} 已接受复核（{reviewer}）")

    # 审计账本（append-only）
    try:
        from app.agents.guardrails import AuditLogger
        AuditLogger.log(
            tool_name="review", action_name="review_accept", risk="medium", agent="review",
            params={"chain_id": content_data.get("chain_id", ""), "message_id": message_id},
            result_preview=content_data.get("verify_summary", ""),
            success=True, session_id=pending_msg.conversation_id,
        )
    except Exception:
        pass

    # 复核结果写入对话（回流通知，含原对话锚点可跳转 + 方案消息定位）
    try:
        _cv_id = content_data.get("conversation_id", "")
        _cv_title = content_data.get("conversation_title", "")
        _cv_mid = content_data.get("message_id", "")
        _cv_ref = ""
        if _cv_id and _cv_title:
            _cv_ref = (f"\n📎 [打开原对话: {_cv_title}](conv://{_cv_id}?mid={_cv_mid})"
                       if _cv_mid else f"\n📎 [打开原对话: {_cv_title}](conv://{_cv_id})")
        elif _cv_title:
            _cv_ref = f"\n📎 原对话: {_cv_title}"
        await repo.create(
            conversation_id=pending_msg.conversation_id, role=MessageRole.ASSISTANT,
            content=(
                f"🟢 复核通过：**{content_data.get('chain_name') or content_data.get('chain_id') or '变更方案'}** "
                f"验证未通过结果已由 {reviewer} 确认接受。\n"
                f"> {content_data.get('verify_summary', '')}"
                + _cv_ref
            ),
            message_type=MessageType.INFO.value,
        )
    except Exception as e:
        log.error(f"[复核] 写入对话失败: {e}")

    # 广播复核完成事件（前端待复核页 + 对话页刷新）
    try:
        from app.services.event_bus import event_bus
        await event_bus.publish("approval_done", {
            "conversation_id": pending_msg.conversation_id,
            "message_id": message_id,
            "action": content_data.get("chain_name", ""),
            "reviewer": reviewer,
            "approved": True,
            "review_type": "accept",
        })
        await event_bus.publish("pending_updated", {"source": "review_accept"})
    except Exception:
        pass

    return {"success": True, "message_id": message_id, "status": updated.status}


@router.post("/{message_id}/review-rollback")
async def review_rollback(
    message_id: str,
    body: ApprovalRequest,
    db: AsyncSession = Depends(get_db),
):
    """复核人选择回滚：触发 {chain_id}_rollback 链撤销变更，标记复核回滚。"""
    from app.repositories.message_repository import MessageRepository
    from app.models.message import MessageType, ConfirmStatus, MessageRole

    repo = MessageRepository(db)
    pending_msg = await repo.get_by_id(message_id)
    if not pending_msg:
        raise HTTPException(status_code=404, detail=f"消息不存在: {message_id}")
    if pending_msg.message_type != MessageType.REVIEW.value:
        raise HTTPException(status_code=400, detail="非复核条目")
    if pending_msg.status != ConfirmStatus.PENDING.value:
        raise HTTPException(status_code=400, detail="该复核已处理")

    content_data = _load_review_entry(pending_msg)
    submitter = content_data.get("submitter_id", "")
    if submitter and body.user_id and submitter == body.user_id:
        raise HTTPException(status_code=400, detail="执行人不能复核自己的变更（责任分离）")

    reviewer = body.user_id or "复核人"
    chain_id = content_data.get("chain_id", "")
    # 触发回滚链（复用 chain_engine 公共回滚逻辑）
    rollback_result = {"triggered": False, "ok": 0, "total": 0, "steps": []}
    if chain_id:
        try:
            from app.core.chain_engine import chain_engine
            rollback_result = await chain_engine._run_rollback_chain(chain_id)
        except Exception as e:
            log.error(f"[复核] 回滚链执行失败 {chain_id}: {e}")
    rolled_back = bool(rollback_result.get("triggered"))
    rb_ok, rb_total = rollback_result.get("ok", 0), rollback_result.get("total", 0)

    updated = await repo.resolve_confirmation(message_id, approved=False, reviewed_by=reviewer)
    if not updated:
        raise HTTPException(status_code=404, detail=f"消息不存在: {message_id}")
    log.info(f"[复核] message_id={message_id} 已回滚（{reviewer}）rolled_back={rolled_back} ({rb_ok}/{rb_total} 步)")

    # 回滚验证结果写回复核条目 content（供已处理列表/卡片展示）
    try:
        rb_vd = rollback_result.get("verify_detail") or []
        if rb_vd:
            content_data["verify_detail"] = rb_vd
            _vok = rollback_result.get("verified")
            content_data["verify_summary"] = (content_data.get("verify_summary", "")
                + f"｜回滚后验证：{'通过' if _vok is True else '未通过'}")
            content_data["rollback_verified"] = _vok
            pending_msg.content = json.dumps(content_data, ensure_ascii=False)
            await db.commit()
    except Exception as e:
        log.error(f"[复核] 回滚验证写回失败: {e}")

    # 审计账本（append-only）
    try:
        from app.agents.guardrails import AuditLogger
        AuditLogger.log(
            tool_name="review", action_name="review_rollback", risk="high", agent="review",
            params={"chain_id": chain_id, "message_id": message_id},
            result_preview=(f"回滚{'成功' if rolled_back else '未触发/失败'} {rb_ok}/{rb_total} 步"
                            + (f"，回滚验证{'通过' if rollback_result.get('verified') is True else '未通过'}"
                               if rollback_result.get('verify_detail') else "")
                            + f": {content_data.get('verify_summary', '')}"),
            success=rolled_back, session_id=pending_msg.conversation_id,
        )
    except Exception:
        pass

    # 复核结果写入对话（回流通知，含回滚步骤详情）
    try:
        roll_msg = (f"🔄 复核回滚：**{content_data.get('chain_name') or chain_id or '变更方案'}** "
                    f"已由 {reviewer} 触发回滚")
        if rolled_back:
            roll_msg += f"\n回滚链执行：**{rb_ok}/{rb_total}** 步成功"
            for st in rollback_result.get("steps", []):
                mark = "✅" if st.get("status") == "success" else "❌"
                desc = st.get("description") or st.get("action_name") or st.get("step_id") or "步骤"
                row = ""
                if st.get("status") == "success" and st.get("rowCount"):
                    if st.get("write"):
                        row = f"，影响 {st.get('rowCount')} 行"
                    else:
                        row = f"，查到 {st.get('rowCount')} 条"
                roll_msg += f"\n- {mark} {desc}{row}"
                if st.get("status") == "error" and st.get("error"):
                    roll_msg += f"：{st['error']}"
            # 回滚后验证：复查回滚链声明的目标状态是否恢复
            rb_vd = rollback_result.get("verify_detail", [])
            if rb_vd:
                for vd in rb_vd:
                    mk = "✅" if vd.get("match") is True else ("❌" if vd.get("match") is False else "⚠")
                    roll_msg += (f"\n回滚后验证：{mk} {vd.get('property')} — "
                                 f"期望 {vd.get('expected', '-')} / 实际 {vd.get('actual', '-')}")
                if rollback_result.get("verified") is False:
                    roll_msg += "\n⚠ 回滚后目标状态未达成，请人工确认。"
        else:
            roll_msg += "，回滚链未找到或执行失败，请人工处理。"
        if body.comment:
            roll_msg += f"\n原因: {body.comment}"
        _cv_id = content_data.get("conversation_id", "")
        _cv_title = content_data.get("conversation_title", "")
        _cv_mid = content_data.get("message_id", "")
        if _cv_id and _cv_title:
            roll_msg += (f"\n📎 [打开原对话: {_cv_title}](conv://{_cv_id}?mid={_cv_mid})"
                         if _cv_mid else f"\n📎 [打开原对话: {_cv_title}](conv://{_cv_id})")
        elif _cv_title:
            roll_msg += f"\n📎 原对话: {_cv_title}"
        await repo.create(
            conversation_id=pending_msg.conversation_id, role=MessageRole.ASSISTANT,
            content=roll_msg, message_type=MessageType.INFO.value,
        )
    except Exception as e:
        log.error(f"[复核] 写入对话失败: {e}")

    # 广播复核完成事件
    try:
        from app.services.event_bus import event_bus
        await event_bus.publish("approval_done", {
            "conversation_id": pending_msg.conversation_id,
            "message_id": message_id,
            "action": content_data.get("chain_name", ""),
            "reviewer": reviewer,
            "approved": False,
            "review_type": "rollback",
        })
        await event_bus.publish("pending_updated", {"source": "review_rollback"})
    except Exception:
        pass

    return {"success": True, "message_id": message_id, "status": updated.status, "rolled_back": rolled_back}
